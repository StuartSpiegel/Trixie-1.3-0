import os

from CommonUtils import read_input
from PopulateStickies import populate_title, populate_color, populate_text, populate_points
from docx import Document

STICKY_TEMPLATE_NAME = "Stickie_template.docx"

# Write result .docx files to this folder
OUTPUT_FOLDER_NAME = "sticky_notes"
OUTPUT_FILE_PREFIX = "stickies"
OUTPUT_KEY_FILE_PREFIX = "color_key"

# There are 31 "run" <w:r> tags in the template Word file.  They are not organized in any discernible way
# Manual inspection revealed the element which corresponds to each run
switcher = {
    0: (0, 0, 6, 21, 26),
    1: (0, 14, 4, 20, 1),
    2: (1, 2, 10, 18, 23),
    3: (0, 15, 9, 19, 24),
    4: (1, 1, 7, 16, 22),
    5: (1, 0, 8, 17, 25)
}


def create_stickies(storyList):
    # Create output folder
    if not os.path.exists(OUTPUT_FOLDER_NAME):
        os.makedirs(OUTPUT_FOLDER_NAME)

    document = Document()
    for i, story in enumerate(storyList):  # TODO: change the iteration logic here to use
        # Make new document at the beginning, or after completing 6 sticky notes
        if i % 6 == 0:
            document = Document(STICKY_TEMPLATE_NAME)

        paragraph, text_run, corner, title_run, color_run = switcher.get(i % 6)
        populate_title(story[4], document.paragraphs[0]._p.r_lst[title_run])
        populate_color(document.paragraphs[0]._p.r_lst[color_run], story[5])
        populate_text(document.paragraphs[paragraph]._p.r_lst[text_run], story[0], story[1], story[2], story[6])
        populate_points(story[3], document.paragraphs[0]._p.r_lst[corner])

        # Save a new file after making 6 sticky notes (or at the end)
        if i % 6 == 5 or i == len(storyList) - 1:
            document.save(OUTPUT_FOLDER_NAME + "/" + OUTPUT_FILE_PREFIX + str(int(i / 6) + 1) + ".docx")


def create_color_key(colorFeatureMap):
    doc = Document()
    count = 0  # Counts how many sticky notes on current page (zero-indexed)
    for color, feature in colorFeatureMap.items():
        # Make new document at the beginning, or after completing 6 sticky notes
        if count % 6 == 0:
            doc = Document(STICKY_TEMPLATE_NAME)

        _, _, _, title_run, color_run = switcher.get(count % 6)
        populate_title(feature, doc.paragraphs[0]._p.r_lst[title_run])
        populate_color(doc.paragraphs[0]._p.r_lst[color_run], color)

        # Save a new file after making 6 sticky notes (or at the end)
        if count % 6 == 5 or count == len(colorFeatureMap) - 1:
            doc.save(OUTPUT_FOLDER_NAME + "/" + OUTPUT_KEY_FILE_PREFIX + str(int(count / 6) + 1) + ".docx")
        count += 1


def print_stories():
    for story in storyList:
        print(story)


storyList, colorFeatureMap = read_input()
# print_stories()
create_stickies(storyList)
create_color_key(colorFeatureMap)
